import logging
import os
from io import BytesIO
from pathlib import Path

from django.conf import settings
from django.core.files.base import ContentFile

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm
except Exception:  # pragma: no cover
    Document = None
    WD_ALIGN_PARAGRAPH = None
    Mm = None


logger = logging.getLogger(__name__)


def _candidate_watermark_paths(raw: str) -> list[Path]:
    if not raw:
        return []

    base_dir = Path(getattr(settings, "BASE_DIR", Path.cwd()))

    raw_path = Path(raw)
    candidates = [
        raw_path,
        base_dir / raw,
        base_dir / "branding" / raw,
        base_dir / "media" / raw,
    ]

    # Убираем дубли, сохраняя порядок
    unique: list[Path] = []
    seen = set()
    for item in candidates:
        key = str(item)
        if key in seen:
            continue
        seen.add(key)
        unique.append(item)

    return unique


def _get_watermark_path():
    raw = getattr(settings, "DOCUMENT_WATERMARK_IMAGE", "") or os.getenv("DOCUMENT_WATERMARK_IMAGE", "")
    if not raw:
        logger.error("DOCUMENT_WATERMARK_IMAGE is empty")
        return None

    candidates = _candidate_watermark_paths(raw)
    for candidate in candidates:
        try:
            if candidate.exists() and candidate.is_file():
                logger.info("Watermark resolved to: %s", candidate)
                return candidate
        except Exception:
            logger.exception("Failed while checking watermark path candidate: %s", candidate)

    logger.error(
        "Watermark file not found. Raw value: %r. Checked: %s",
        raw,
        [str(p) for p in candidates],
    )
    return None


def _iter_unique_footers(section):
    seen = set()
    for footer in (
        getattr(section, "footer", None),
        getattr(section, "first_page_footer", None),
        getattr(section, "even_page_footer", None),
    ):
        if footer is None:
            continue

        footer_id = id(footer)
        if footer_id in seen:
            continue

        seen.add(footer_id)
        yield footer


def _append_watermark_to_footer(footer, watermark_path):
    if getattr(footer, "is_linked_to_previous", False):
        footer.is_linked_to_previous = False

    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = paragraph.add_run()
    run.add_picture(str(watermark_path), width=Mm(35))


def build_approved_document(generated_document):
    source_field = getattr(generated_document, "generated_file", None)
    if not source_field:
        logger.error("Generated document %s has no generated_file", getattr(generated_document, "id", None))
        return None

    source_name = getattr(source_field, "name", "") or ""
    source_path_raw = getattr(source_field, "path", "") or ""
    source_path = Path(source_path_raw) if source_path_raw else None

    if not source_name.lower().endswith(".docx"):
        logger.error(
            "Generated document %s is not a .docx file: %s",
            getattr(generated_document, "id", None),
            source_name,
        )
        return None

    if not source_path or not source_path.exists():
        logger.error(
            "Generated document file path does not exist for document %s: %s",
            getattr(generated_document, "id", None),
            source_path_raw,
        )
        return None

    if Document is None or WD_ALIGN_PARAGRAPH is None or Mm is None:
        logger.error("python-docx imports are unavailable")
        return None

    watermark_path = _get_watermark_path()
    if watermark_path is None:
        logger.error("Watermark path could not be resolved")
        return None

    try:
        doc = Document(str(source_path))
    except Exception:
        logger.exception(
            "Failed to open generated .docx for document %s: %s",
            getattr(generated_document, "id", None),
            source_path,
        )
        return None

    try:
        added = 0
        for section in doc.sections:
            for footer in _iter_unique_footers(section):
                _append_watermark_to_footer(footer, watermark_path)
                added += 1

        if added == 0:
            logger.error(
                "No footer targets found while applying watermark for document %s",
                getattr(generated_document, "id", None),
            )
            return None

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        approved_name = f"generated_documents/approved/approved_{Path(source_name).name}"
        logger.info(
            "Approved document built successfully for document %s with watermark %s",
            getattr(generated_document, "id", None),
            watermark_path,
        )
        return ContentFile(buffer.read(), name=approved_name)

    except Exception:
        logger.exception(
            "Failed while applying watermark for document %s using %s",
            getattr(generated_document, "id", None),
            watermark_path,
        )
        return None