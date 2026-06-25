import io
import shutil
import subprocess
import tempfile
from pathlib import Path
from uuid import uuid4

import fitz
from django.conf import settings
from django.core.files.base import ContentFile
from django.utils.text import get_valid_filename
from django.db import models, transaction
from django.template.defaultfilters import slugify
from django.utils import timezone
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
from docxtpl import DocxTemplate

from apps.core.models import ActiveModel, OrderedModel, TimeStampedModel
from apps.crm.models import Application, Client
from apps.finance.models import Deal
from apps.organizations.models import Company, Office


def template_upload_path(instance, filename):
    return f'erp/documents/templates/{instance.company_id or "global"}/{filename}'


def generated_upload_path(instance, filename):
    return f'erp/documents/generated/{instance.company_id or "global"}/{filename}'


def approved_upload_path(instance, filename):
    return f'erp/documents/approved/{instance.company_id or "global"}/{filename}'


def stamp_preview_upload_path(instance, filename):
    return f'erp/documents/stamp_previews/{instance.company_id or "global"}/{filename}'


def stamp_upload_path(instance, filename):
    return f'erp/documents/stamps/{instance.company_id or "global"}/{filename}'


def safe_text(value):
    if value is None:
        return ''
    return str(value)


def get_cyrillic_font_path():
    candidates = [
        getattr(settings, 'PDF_CYRILLIC_FONT_PATH', ''),
        '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
        '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
        '/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf',
        'C:/Windows/Fonts/arial.ttf',
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return str(candidate)
    return None


def fitz_text_kwargs(font_size=None):
    font_path = get_cyrillic_font_path()
    kwargs = {'fontname': 'helv'}
    if font_path:
        kwargs = {'fontname': 'DejaVuSans', 'fontfile': font_path}
    if font_size is not None:
        kwargs['fontsize'] = font_size
    return kwargs


def user_display_name(user):
    if not user:
        return ''
    return user.get_full_name() or getattr(user, 'email', '') or safe_text(user)


def clean_download_name(value):
    value = safe_text(value).strip().replace('/', '-').replace('\\', '-')
    value = ' '.join(value.split())
    return (get_valid_filename(value) or 'document')[:140].strip(' ._-') or 'document'


def safe_file_slug(value, fallback='document', max_length=80):
    slug = slugify(safe_text(value)).strip('-_')[:max_length].strip('-_')
    return slug or fallback


def safe_document_title(template_name='', client_name='', title='', max_length=255):
    template_name = safe_text(template_name).strip() or 'Документ'
    client_name = safe_text(client_name).strip()
    title = safe_text(title).strip()

    if not title or 'без клиента' in title.lower():
        title = f'{template_name} - {client_name}' if client_name else template_name
    elif client_name and client_name.lower() not in title.lower():
        title = f'{title} - {client_name}'

    title = ' '.join(title.split())
    if len(title) > max_length:
        title = title[:max_length].rstrip(' ._-')
    return title or template_name[:max_length]


class DocumentTemplate(TimeStampedModel, ActiveModel):
    company = models.ForeignKey(
        Company,
        verbose_name='Компания',
        on_delete=models.PROTECT,
        related_name='erp_document_templates',
        null=True,
        blank=True,
    )
    name = models.CharField('Название шаблона', max_length=255, db_index=True)
    code = models.SlugField('Код шаблона', max_length=100, db_index=True)
    document_type = models.CharField('Тип документа', max_length=100, blank=True)
    description = models.TextField('Описание', blank=True)
    file = models.FileField('DOCX-файл шаблона', upload_to=template_upload_path)
    requires_approval = models.BooleanField('Требуется подтверждение администратора', default=True)
    allow_without_stamp = models.BooleanField('Разрешить скачать DOCX без печати', default=True)
    allow_with_stamp = models.BooleanField('Разрешить скачать PDF с печатью', default=True)
    jinja_variables = models.JSONField('Jinja-переменные', default=list, blank=True)
    stamp_settings = models.JSONField('Технические настройки печати', default=dict, blank=True)
    watermark_settings = models.JSONField('Технические настройки водяного знака', default=dict, blank=True)
    created_by = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        verbose_name='Кем создан',
        on_delete=models.SET_NULL,
        related_name='erp_document_templates_created',
        null=True,
        blank=True,
    )

    class Meta:
        verbose_name = 'Шаблон документа'
        verbose_name_plural = 'Шаблоны документов'
        ordering = ['company__name', 'name']
        unique_together = [('company', 'code')]
        indexes = [
            models.Index(fields=['company', 'is_active']),
            models.Index(fields=['code']),
            models.Index(fields=['name']),
        ]

    def __str__(self):
        return self.name

    def save(self, *args, **kwargs):
        if not self.code:
            self.code = slugify(self.name)[:100] or f'template-{self.pk or "new"}'
        super().save(*args, **kwargs)


class DocumentTemplateField(TimeStampedModel, OrderedModel):
    SOURCE_CLIENT = 'client'
    SOURCE_APPLICATION = 'application'
    SOURCE_DEAL = 'deal'
    SOURCE_MANAGER = 'manager'
    SOURCE_COMPANY = 'company'
    SOURCE_OFFICE = 'office'
    SOURCE_CUSTOM = 'custom'
    SOURCE_CHOICES = (
        (SOURCE_CLIENT, 'Клиент'),
        (SOURCE_APPLICATION, 'Заявка'),
        (SOURCE_DEAL, 'Сделка'),
        (SOURCE_MANAGER, 'Менеджер'),
        (SOURCE_COMPANY, 'Компания'),
        (SOURCE_OFFICE, 'Офис'),
        (SOURCE_CUSTOM, 'Свое значение'),
    )

    FIELD_TYPE_TEXT = 'text'
    FIELD_TYPE_TEXTAREA = 'textarea'
    FIELD_TYPE_NUMBER = 'number'
    FIELD_TYPE_DATE = 'date'
    FIELD_TYPE_BOOLEAN = 'boolean'
    FIELD_TYPE_SELECT = 'select'
    FIELD_TYPE_CHOICES = (
        (FIELD_TYPE_TEXT, 'Текст'),
        (FIELD_TYPE_TEXTAREA, 'Большой текст'),
        (FIELD_TYPE_NUMBER, 'Число'),
        (FIELD_TYPE_DATE, 'Дата'),
        (FIELD_TYPE_BOOLEAN, 'Да/нет'),
        (FIELD_TYPE_SELECT, 'Выбор'),
    )

    template = models.ForeignKey(
        DocumentTemplate,
        verbose_name='Шаблон',
        on_delete=models.CASCADE,
        related_name='fields',
    )
    key = models.SlugField('Ключ поля', max_length=100)
    jinja_key = models.CharField('Jinja-ключ', max_length=150, blank=True)
    data_source = models.CharField('Источник данных', max_length=32, choices=SOURCE_CHOICES, default=SOURCE_CUSTOM, db_index=True)
    label = models.CharField('Название поля', max_length=255)
    field_type = models.CharField('Тип поля', max_length=32, choices=FIELD_TYPE_CHOICES, default=FIELD_TYPE_TEXT)
    default_value = models.CharField('Значение по умолчанию', max_length=255, blank=True)
    options = models.JSONField('Варианты выбора', default=list, blank=True)
    is_required = models.BooleanField('Обязательное поле', default=True)
    help_text = models.CharField('Подсказка', max_length=255, blank=True)

    class Meta:
        verbose_name = 'Поле шаблона документа'
        verbose_name_plural = 'Поля шаблонов документов'
        ordering = ['sort_order', 'label']
        unique_together = [('template', 'key')]
        indexes = [
            models.Index(fields=['template', 'sort_order']),
        ]

    def __str__(self):
        return f'{self.template}: {self.label}'

    def save(self, *args, **kwargs):
        if not self.jinja_key:
            self.jinja_key = self.key
        super().save(*args, **kwargs)


class GeneratedDocument(TimeStampedModel):
    STATUS_DRAFT = 'draft'
    STATUS_GENERATED = 'generated'
    STATUS_PENDING = 'pending'
    STATUS_APPROVED = 'approved'
    STATUS_REJECTED = 'rejected'
    STATUS_ERROR = 'error'
    STATUS_CHOICES = (
        (STATUS_DRAFT, 'Черновик'),
        (STATUS_GENERATED, 'Сгенерирован'),
        (STATUS_PENDING, 'Ожидает подтверждения'),
        (STATUS_APPROVED, 'Одобрен'),
        (STATUS_REJECTED, 'Отклонён'),
        (STATUS_ERROR, 'Ошибка генерации'),
    )

    company = models.ForeignKey(Company, verbose_name='Компания', on_delete=models.PROTECT, related_name='erp_generated_documents')
    office = models.ForeignKey(
        Office,
        verbose_name='Офис',
        on_delete=models.SET_NULL,
        related_name='erp_generated_documents',
        null=True,
        blank=True,
    )
    template = models.ForeignKey(DocumentTemplate, verbose_name='Шаблон', on_delete=models.PROTECT, related_name='generated_documents')
    client = models.ForeignKey(
        Client,
        verbose_name='Клиент',
        on_delete=models.PROTECT,
        related_name='erp_generated_documents',
        null=True,
        blank=True,
    )
    application = models.ForeignKey(
        Application,
        verbose_name='Заявка',
        on_delete=models.SET_NULL,
        related_name='erp_generated_documents',
        null=True,
        blank=True,
    )
    deal = models.ForeignKey(
        Deal,
        verbose_name='Сделка',
        on_delete=models.SET_NULL,
        related_name='erp_generated_documents',
        null=True,
        blank=True,
    )
    manager = models.ForeignKey(settings.AUTH_USER_MODEL, verbose_name='Менеджер', on_delete=models.PROTECT, related_name='erp_generated_documents')
    title = models.CharField('Название документа', max_length=255, blank=True)
    context_data = models.JSONField('Данные для подстановки', default=dict, blank=True)
    status = models.CharField('Статус', max_length=32, choices=STATUS_CHOICES, default=STATUS_DRAFT, db_index=True)
    generated_file = models.FileField('DOCX без печати', upload_to=generated_upload_path, null=True, blank=True)
    approved_file = models.FileField('PDF с печатью / одобренный файл', upload_to=approved_upload_path, null=True, blank=True)
    stamp_preview_file = models.FileField('Предпросмотр PDF с печатью', upload_to=stamp_preview_upload_path, null=True, blank=True)
    stamp_preview_options = models.JSONField('Настройки предпросмотра печати', default=dict, blank=True)
    stamp_preview_generated_at = models.DateTimeField('Предпросмотр печати создан', null=True, blank=True)
    stamp_preview_generated_by = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        verbose_name='Кем создан предпросмотр печати',
        on_delete=models.SET_NULL,
        related_name='erp_documents_stamp_previews',
        null=True,
        blank=True,
    )
    generation_error = models.TextField('Ошибка генерации', blank=True)
    submitted_at = models.DateTimeField('Отправлен на подтверждение', null=True, blank=True)
    generated_at = models.DateTimeField('Сгенерирован', null=True, blank=True)
    approved_by = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        verbose_name='Кем одобрен',
        on_delete=models.SET_NULL,
        related_name='erp_documents_approved',
        null=True,
        blank=True,
    )
    approved_at = models.DateTimeField('Дата одобрения', null=True, blank=True)

    class Meta:
        verbose_name = 'Сгенерированный документ'
        verbose_name_plural = 'Сгенерированные документы'
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['company', 'office', 'status']),
            models.Index(fields=['manager', 'status']),
            models.Index(fields=['client', 'status']),
            models.Index(fields=['created_at']),
        ]

    def __str__(self):
        return self.display_title

    @property
    def resolved_client(self):
        client = getattr(self, 'client', None)
        if client:
            return client
        application = getattr(self, 'application', None)
        if application and getattr(application, 'client', None):
            return application.client
        deal = getattr(self, 'deal', None)
        if deal and getattr(deal, 'client', None):
            return deal.client
        return None

    @property
    def display_title(self):
        title = (self.title or '').strip()
        client = self.resolved_client
        client_name = (getattr(client, 'full_name', '') or '').strip()
        template = getattr(self, 'template', None)
        template_name = template.name if template else 'Документ'

        if title or client_name:
            return safe_document_title(template_name, client_name, title)
        return f'{template_name} #{self.pk or ""}'.strip()

    def download_filename(self, file_type='original'):
        if file_type in {'approved', 'pdf', DocumentDownloadLog.FILE_TYPE_APPROVED}:
            extension = self.approved_file_extension or 'pdf'
            suffix = 'с печатью' if extension == 'pdf' else 'одобренный документ'
        else:
            extension = 'docx'
            suffix = 'без печати'
        return f'{clean_download_name(f"{self.display_title} - {suffix}")}.{extension}'

    @property
    def can_download_original(self):
        return bool(self.template.allow_without_stamp and self.generated_file and self.status in {
            self.STATUS_GENERATED,
            self.STATUS_PENDING,
            self.STATUS_APPROVED,
            self.STATUS_REJECTED,
        })

    @property
    def can_download_approved(self):
        return bool(self.approved_file and self.status == self.STATUS_APPROVED)

    @property
    def approved_file_extension(self):
        if not self.approved_file:
            return ''
        name = str(getattr(self.approved_file, 'name', '') or '')
        return Path(name).suffix.lower().lstrip('.') or 'pdf'

    @property
    def approved_file_is_pdf(self):
        return self.approved_file_extension == 'pdf'

    @property
    def can_preview_approved(self):
        return bool(self.can_download_approved and self.approved_file_is_pdf)

    @property
    def approved_download_label(self):
        if self.approved_file_is_pdf:
            return 'Скачать PDF с печатью'
        return 'Скачать одобренный DOCX'

    @property
    def has_stamp_preview(self):
        return bool(self.stamp_preview_file)

    @property
    def can_download(self):
        return self.can_download_original or self.can_download_approved

    @property
    def rejection_reason(self):
        try:
            return self.approval.rejection_reason
        except Exception:
            return ''

    def build_context(self):
        context = {}
        client = self.client or (self.application.client if self.application_id else None) or (self.deal.client if self.deal_id else None)
        application = self.application or (self.deal.application if self.deal_id else None)
        deal = self.deal
        company = self.company
        office = self.office
        manager = self.manager

        if company:
            context['company'] = {
                'name': company.name,
                'legal_name': company.legal_name,
                'phone': company.phone,
                'email': company.email,
                'address': company.address,
                'city': company.city,
                'country': company.country,
            }
            context.update({
                'company_name': company.name,
                'company_legal_name': company.legal_name,
                'company_phone': company.phone,
                'company_email': company.email,
                'company_address': company.address,
            })

        if office:
            context['office'] = {
                'name': office.name,
                'city': office.city,
                'address': office.address,
                'phone': office.phone,
                'email': office.email,
            }
            context.update({
                'office_name': office.name,
                'office_city': office.city,
                'office_address': office.address,
                'office_phone': office.phone,
                'office_email': office.email,
            })

        if manager:
            context['manager'] = {
                'first_name': getattr(manager, 'first_name', ''),
                'last_name': getattr(manager, 'last_name', ''),
                'middle_name': getattr(manager, 'middle_name', ''),
                'name': user_display_name(manager),
                'email': getattr(manager, 'email', ''),
                'phone': getattr(manager, 'phone', ''),
            }
            context.update({
                'manager_name': user_display_name(manager),
                'manager_email': getattr(manager, 'email', ''),
                'manager_phone': getattr(manager, 'phone', ''),
            })

        if client:
            context['client'] = {
                'full_name': client.full_name,
                'phone': client.phone,
                'email': client.email or '',
                'dob': safe_text(client.dob),
                'citizenship': client.citizenship,
                'city': client.city,
                'address': client.address,
                'address_registration': client.address_registration,
                'passport_local_num': client.passport_local_num,
                'passport_inter_num': client.passport_inter_num,
                'passport_issued_by': client.passport_issued_by,
                'passport_issued_date': safe_text(client.passport_issued_date),
                'passport_valid_until': safe_text(getattr(client, 'passport_valid_until', '')),
                'passport_birth_place': getattr(client, 'passport_birth_place', ''),
                'direction': getattr(client, 'direction', ''),
                'interested_country': getattr(client, 'interested_country', ''),
                'interested_university': getattr(client, 'interested_university', ''),
                'interested_program': getattr(client, 'interested_program', ''),
            }
            context.update({
                'client_full_name': client.full_name,
                'client_phone': client.phone,
                'client_email': client.email or '',
                'client_dob': safe_text(client.dob),
                'client_citizenship': client.citizenship,
                'client_city': client.city,
                'client_address': client.address,
                'client_address_registration': client.address_registration,
                'client_passport_local_num': client.passport_local_num,
                'client_passport_inter_num': client.passport_inter_num,
                'client_passport_issued_by': client.passport_issued_by,
                'client_passport_issued_date': safe_text(client.passport_issued_date),
                'client_passport_valid_until': safe_text(getattr(client, 'passport_valid_until', '')),
                'client_passport_birth_place': getattr(client, 'passport_birth_place', ''),
            })

        if application:
            context['application'] = {
                'university_name': application.university_name,
                'program_name': application.program_name,
                'country': application.country,
                'degree': application.degree,
                'language': application.language,
                'intake': application.intake,
                'status': application.get_status_display(),
            }
            context.update({
                'application_university_name': application.university_name,
                'application_program_name': application.program_name,
                'application_country': application.country,
                'application_degree': application.degree,
                'application_language': application.language,
                'application_intake': application.intake,
                'application_status': application.get_status_display(),
            })

        if deal:
            context['deal'] = {
                'title': deal.title,
                'type': deal.get_deal_type_display(),
                'university_name': deal.university_name,
                'program_name': deal.program_name,
                'service_title': deal.service.title if deal.service_id else '',
                'price_client': safe_text(deal.price_client),
                'currency': deal.currency.code if deal.currency_id else '',
                'total_to_pay_usd': safe_text(deal.total_to_pay_usd),
                'paid_amount_usd': safe_text(deal.paid_amount_usd),
                'payment_status': deal.get_payment_status_display(),
            }
            context.update({
                'deal_title': deal.title,
                'deal_type': deal.get_deal_type_display(),
                'deal_university_name': deal.university_name,
                'deal_program_name': deal.program_name,
                'deal_price_client': safe_text(deal.price_client),
                'deal_total_to_pay_usd': safe_text(deal.total_to_pay_usd),
                'deal_paid_amount_usd': safe_text(deal.paid_amount_usd),
                'deal_payment_status': deal.get_payment_status_display(),
            })

        for field in self.template.fields.all():
            if field.default_value and field.key not in context:
                context[field.key] = field.default_value

        context.update(self.context_data or {})
        return context

    def validate_required_fields(self, context):
        missing = []
        for field in self.template.fields.filter(is_required=True):
            value = context.get(field.key)
            if value in (None, '', []):
                missing.append(field.key)
        return missing

    def generate_file(self):
        if not self.template.file:
            raise ValueError('Template file is required.')

        context = self.build_context()
        missing = self.validate_required_fields(context)
        if missing:
            raise ValueError(f'Missing required template fields: {", ".join(missing)}')

        template_path = Path(self.template.file.path)
        doc = DocxTemplate(template_path)
        doc.render(context)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        if not self.title or 'без клиента' in self.title.lower():
            client = self.resolved_client
            self.title = safe_document_title(
                self.template.name,
                getattr(client, 'full_name', '') if client else '',
                self.title,
            )

        base = safe_file_slug(self.display_title, 'document')
        filename = f'{base}-{uuid4().hex[:10]}.docx'

        if self.generated_file:
            self.generated_file.delete(save=False)
        if self.stamp_preview_file:
            self.stamp_preview_file.delete(save=False)
            self.stamp_preview_file = None
        self.stamp_preview_options = {}
        self.stamp_preview_generated_at = None
        self.stamp_preview_generated_by = None

        self.generated_file.save(filename, ContentFile(buffer.getvalue()), save=False)
        self.status = self.STATUS_GENERATED
        self.generation_error = ''
        self.generated_at = timezone.now()
        self.save(update_fields=[
            'generated_file',
            'stamp_preview_file',
            'stamp_preview_options',
            'stamp_preview_generated_at',
            'stamp_preview_generated_by',
            'status',
            'generation_error',
            'generated_at',
            'title',
            'updated_at',
        ])
        return self

    def submit_for_approval(self, user=None, comment=''):
        if not self.generated_file:
            self.generate_file()

        with transaction.atomic():
            approval, _ = DocumentApproval.objects.get_or_create(document=self)
            approval.status = DocumentApproval.STATUS_PENDING
            approval.approval_type = DocumentApproval.TYPE_NOT_SELECTED
            approval.comment = comment or ''
            approval.rejection_reason = ''
            approval.reviewed_by = None
            approval.reviewed_at = None
            approval.save()

            self.status = self.STATUS_PENDING
            self.submitted_at = timezone.now()
            self.approved_by = None
            self.approved_at = None
            if self.approved_file:
                self.approved_file.delete(save=False)
                self.approved_file = None
            if self.stamp_preview_file:
                self.stamp_preview_file.delete(save=False)
                self.stamp_preview_file = None
            self.stamp_preview_options = {}
            self.stamp_preview_generated_at = None
            self.stamp_preview_generated_by = None
            self.save(update_fields=[
                'status',
                'submitted_at',
                'approved_by',
                'approved_at',
                'approved_file',
                'stamp_preview_file',
                'stamp_preview_options',
                'stamp_preview_generated_at',
                'stamp_preview_generated_by',
                'updated_at',
            ])
        return self

    def generate_stamp_preview(self, user=None, stamp_options=None):
        if not self.generated_file:
            raise ValueError('Generated file is required before stamp preview.')
        if self.status == self.STATUS_APPROVED:
            raise ValueError('Подтверждённый документ нельзя перегенерировать. Создайте новый документ.')
        if not self.template.allow_with_stamp:
            raise ValueError('Этот шаблон не разрешает PDF с электронной печатью.')

        stamp_options = stamp_options or {'stamp_mode': 'executor'}
        try:
            preview_file = build_fast_stamp_preview_file(self, stamp_options=stamp_options)
        except Exception as exc:
            self.generation_error = str(exc)
            self.save(update_fields=['generation_error', 'updated_at'])
            raise ValueError(str(exc)) from exc

        if self.stamp_preview_file:
            self.stamp_preview_file.delete(save=False)
        self.stamp_preview_file.save(preview_file[0], preview_file[1], save=False)
        self.stamp_preview_options = stamp_options
        self.stamp_preview_generated_at = timezone.now()
        self.stamp_preview_generated_by = user
        self.generation_error = ''
        self.save(update_fields=[
            'stamp_preview_file',
            'stamp_preview_options',
            'stamp_preview_generated_at',
            'stamp_preview_generated_by',
            'generation_error',
            'updated_at',
        ])
        return self

    def mark_approval_record_approved(self, user, approval_type, comment='', reviewed_at=None):
        reviewed_at = reviewed_at or timezone.now()
        approval, _ = DocumentApproval.objects.get_or_create(document=self)
        approval.status = DocumentApproval.STATUS_APPROVED
        approval.approval_type = approval_type
        approval.comment = comment or approval.comment or ''
        approval.rejection_reason = ''
        approval.reviewed_by = user or approval.reviewed_by
        approval.reviewed_at = reviewed_at
        approval.save()
        return approval

    def already_approved_result(self, user=None, approval_type=None, comment=''):
        if self.status != self.STATUS_APPROVED or not self.approved_file:
            return None
        update_fields = []
        if self.generation_error:
            self.generation_error = ''
            update_fields.append('generation_error')
        if not self.approved_by and user:
            self.approved_by = user
            update_fields.append('approved_by')
        if not self.approved_at:
            self.approved_at = timezone.now()
            update_fields.append('approved_at')
        if update_fields:
            update_fields.append('updated_at')
            self.save(update_fields=update_fields)
        actual_approval_type = (
            DocumentApproval.TYPE_WITH_STAMP
            if self.approved_file_is_pdf
            else DocumentApproval.TYPE_WITHOUT_STAMP
        )
        self.mark_approval_record_approved(user, actual_approval_type, comment=comment, reviewed_at=self.approved_at)
        return self

    def approve_stamp_preview(self, user, comment=''):
        approved = self.already_approved_result(user, DocumentApproval.TYPE_WITH_STAMP, comment=comment)
        if approved:
            return approved
        if not self.stamp_preview_file:
            raise ValueError('Сначала сгенерируйте предпросмотр PDF с печатью и проверьте его.')
        if not self.template.allow_with_stamp:
            raise ValueError('Этот шаблон не разрешает PDF с электронной печатью.')

        stamp_options = self.stamp_preview_options or {'stamp_mode': 'executor'}
        try:
            approved_file = build_approved_document_file(self, with_stamp=True, stamp_options=stamp_options)
        except Exception as exc:
            self.generation_error = str(exc)
            self.save(update_fields=['generation_error', 'updated_at'])
            raise ValueError(str(exc)) from exc

        with transaction.atomic():
            if self.approved_file:
                self.approved_file.delete(save=False)
            self.approved_file.save(approved_file[0], approved_file[1], save=False)
            self.status = self.STATUS_APPROVED
            self.approved_by = user
            self.approved_at = timezone.now()
            self.generation_error = ''
            stored_context = self.context_data or {}
            stored_context['last_stamp_options'] = stamp_options
            stored_context['approved_from_stamp_preview'] = False
            stored_context['approved_from_original_docx'] = True
            self.context_data = stored_context
            self.save(update_fields=[
                'approved_file',
                'status',
                'approved_by',
                'approved_at',
                'generation_error',
                'context_data',
                'updated_at',
            ])

            self.mark_approval_record_approved(
                user,
                DocumentApproval.TYPE_WITH_STAMP,
                comment=comment,
                reviewed_at=self.approved_at,
            )
        return self

    def approve(self, user, with_stamp=False, comment='', stamp_options=None):
        if not self.generated_file:
            raise ValueError('Generated file is required before approval.')
        if with_stamp and not self.template.allow_with_stamp:
            raise ValueError('Этот шаблон не разрешает PDF с электронной печатью.')
        if not with_stamp and not self.template.allow_without_stamp:
            raise ValueError('Этот шаблон не разрешает подтверждение без печати.')

        approval_type = DocumentApproval.TYPE_WITH_STAMP if with_stamp else DocumentApproval.TYPE_WITHOUT_STAMP
        approved = self.already_approved_result(user, approval_type, comment=comment)
        if approved:
            return approved
        try:
            approved_file = build_approved_document_file(self, with_stamp=with_stamp, stamp_options=stamp_options)
        except Exception as exc:
            self.generation_error = str(exc)
            self.save(update_fields=['generation_error', 'updated_at'])
            raise ValueError(str(exc)) from exc

        with transaction.atomic():
            if self.approved_file:
                self.approved_file.delete(save=False)
            self.approved_file.save(approved_file[0], approved_file[1], save=False)
            self.status = self.STATUS_APPROVED
            self.approved_by = user
            self.approved_at = timezone.now()
            self.generation_error = ''
            if with_stamp:
                stored_context = self.context_data or {}
                stored_context['last_stamp_options'] = stamp_options or {'stamp_mode': 'executor'}
                self.context_data = stored_context
            self.save(update_fields=['approved_file', 'status', 'approved_by', 'approved_at', 'generation_error', 'context_data', 'updated_at'])

            self.mark_approval_record_approved(user, approval_type, comment=comment, reviewed_at=self.approved_at)
        return self

    def reject(self, user, reason=''):
        with transaction.atomic():
            self.status = self.STATUS_REJECTED
            self.save(update_fields=['status', 'updated_at'])

            approval, _ = DocumentApproval.objects.get_or_create(document=self)
            approval.status = DocumentApproval.STATUS_REJECTED
            approval.rejection_reason = reason or ''
            approval.reviewed_by = user
            approval.reviewed_at = timezone.now()
            approval.save()
        return self


class DocumentApproval(TimeStampedModel):
    STATUS_PENDING = 'pending'
    STATUS_APPROVED = 'approved'
    STATUS_REJECTED = 'rejected'
    STATUS_CHOICES = (
        (STATUS_PENDING, 'Ожидает проверки'),
        (STATUS_APPROVED, 'Одобрено'),
        (STATUS_REJECTED, 'Отклонено'),
    )

    TYPE_NOT_SELECTED = 'not_selected'
    TYPE_WITHOUT_STAMP = 'without_stamp'
    TYPE_WITH_STAMP = 'with_stamp'
    TYPE_CHOICES = (
        (TYPE_NOT_SELECTED, 'Не выбрано'),
        (TYPE_WITHOUT_STAMP, 'Одобрить без печати'),
        (TYPE_WITH_STAMP, 'Одобрить с электронной печатью'),
    )

    document = models.OneToOneField(GeneratedDocument, verbose_name='Документ', on_delete=models.CASCADE, related_name='approval')
    status = models.CharField('Статус', max_length=32, choices=STATUS_CHOICES, default=STATUS_PENDING, db_index=True)
    approval_type = models.CharField('Тип подтверждения', max_length=32, choices=TYPE_CHOICES, default=TYPE_NOT_SELECTED)
    comment = models.TextField('Комментарий', blank=True)
    rejection_reason = models.TextField('Причина отклонения', blank=True)
    reviewed_by = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        verbose_name='Кем проверено',
        on_delete=models.SET_NULL,
        related_name='erp_document_approvals',
        null=True,
        blank=True,
    )
    reviewed_at = models.DateTimeField('Дата проверки', null=True, blank=True)

    class Meta:
        verbose_name = 'Подтверждение документа'
        verbose_name_plural = 'Подтверждения документов'
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['status', 'reviewed_at']),
        ]

    def __str__(self):
        return f'{self.document} - {self.status}'


class StampRule(TimeStampedModel, ActiveModel, OrderedModel):
    POSITION_BOTTOM_LEFT = 'bottom_left'
    POSITION_BOTTOM_RIGHT = 'bottom_right'
    POSITION_TOP_LEFT = 'top_left'
    POSITION_TOP_RIGHT = 'top_right'
    POSITION_CENTER = 'center'
    POSITION_CUSTOM = 'custom'
    POSITION_CHOICES = (
        (POSITION_BOTTOM_LEFT, 'Снизу слева'),
        (POSITION_BOTTOM_RIGHT, 'Снизу справа'),
        (POSITION_TOP_LEFT, 'Сверху слева'),
        (POSITION_TOP_RIGHT, 'Сверху справа'),
        (POSITION_CENTER, 'По центру'),
        (POSITION_CUSTOM, 'Своя позиция'),
    )

    company = models.ForeignKey(
        Company,
        verbose_name='Компания',
        on_delete=models.CASCADE,
        related_name='erp_document_stamp_rules',
        null=True,
        blank=True,
    )
    office = models.ForeignKey(
        Office,
        verbose_name='Офис',
        on_delete=models.CASCADE,
        related_name='erp_document_stamp_rules',
        null=True,
        blank=True,
    )
    template = models.ForeignKey(
        DocumentTemplate,
        verbose_name='Шаблон',
        on_delete=models.CASCADE,
        related_name='stamp_rules',
        null=True,
        blank=True,
    )
    name = models.CharField('Название правила', max_length=150)
    stamp_image = models.ImageField('Файл электронной печати', upload_to=stamp_upload_path)
    width_mm = models.PositiveIntegerField('Ширина печати, мм', default=40)
    height_mm = models.PositiveIntegerField('Высота печати, мм', default=40)
    position = models.CharField('Позиция печати', max_length=32, choices=POSITION_CHOICES, default=POSITION_BOTTOM_LEFT)
    x_mm = models.DecimalField('Координата X, мм', max_digits=8, decimal_places=2, null=True, blank=True)
    y_mm = models.DecimalField('Координата Y, мм', max_digits=8, decimal_places=2, null=True, blank=True)
    opacity = models.DecimalField('Прозрачность печати', max_digits=4, decimal_places=2, default=1)
    watermark_enabled = models.BooleanField('Включить водяной знак', default=False)
    watermark_text = models.CharField('Текст водяного знака', max_length=255, blank=True)
    watermark_image = models.ImageField('Изображение водяного знака', upload_to=stamp_upload_path, null=True, blank=True)
    watermark_position = models.CharField('Позиция водяного знака', max_length=32, choices=POSITION_CHOICES, default=POSITION_CENTER)
    watermark_width_mm = models.PositiveIntegerField('Ширина водяного знака, мм', default=160)
    watermark_height_mm = models.PositiveIntegerField('Высота водяного знака, мм', default=60)
    watermark_opacity = models.DecimalField('Прозрачность водяного знака', max_digits=4, decimal_places=2, default=0.15)

    class Meta:
        verbose_name = 'Правило электронной печати'
        verbose_name_plural = 'Правила электронной печати'
        ordering = ['sort_order', 'name']
        indexes = [
            models.Index(fields=['company', 'office', 'is_active']),
            models.Index(fields=['template', 'is_active']),
        ]

    def __str__(self):
        return self.name


class DocumentDownloadLog(models.Model):
    FILE_TYPE_ORIGINAL = 'original'
    FILE_TYPE_APPROVED = 'approved'
    FILE_TYPE_CHOICES = (
        (FILE_TYPE_ORIGINAL, 'DOCX без печати'),
        (FILE_TYPE_APPROVED, 'PDF с печатью'),
    )

    document = models.ForeignKey(GeneratedDocument, verbose_name='Документ', on_delete=models.CASCADE, related_name='download_logs')
    user = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        verbose_name='Пользователь',
        on_delete=models.SET_NULL,
        related_name='erp_document_download_logs',
        null=True,
        blank=True,
    )
    file_type = models.CharField('Тип файла', max_length=32, choices=FILE_TYPE_CHOICES)
    ip_address = models.GenericIPAddressField('IP-адрес', null=True, blank=True)
    user_agent = models.TextField('User-Agent', blank=True)
    created_at = models.DateTimeField('Дата скачивания', auto_now_add=True, db_index=True)

    class Meta:
        verbose_name = 'Лог скачивания документа'
        verbose_name_plural = 'Логи скачивания документов'
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['document', 'file_type']),
            models.Index(fields=['user', 'created_at']),
        ]

    def __str__(self):
        return f'{self.document} - {self.file_type}'


def find_stamp_rule(document):
    qs = StampRule.objects.filter(is_active=True)
    candidates = [
        {'template': document.template, 'office': document.office, 'company': document.company},
        {'template': document.template, 'office__isnull': True, 'company': document.company},
        {'template': document.template, 'office__isnull': True, 'company__isnull': True},
        {'template__isnull': True, 'office': document.office, 'company': document.company},
        {'template__isnull': True, 'office__isnull': True, 'company': document.company},
        {'template__isnull': True, 'office__isnull': True, 'company__isnull': True},
    ]
    for filters in candidates:
        rule = qs.filter(**filters).order_by('sort_order', 'id').first()
        if rule:
            return rule
    return None


def copy_generated_file(document):
    document.generated_file.open('rb')
    try:
        content = document.generated_file.read()
    finally:
        document.generated_file.close()

    base = safe_file_slug(document.display_title, 'approved-document')
    return f'{base}-approved-{uuid4().hex[:10]}.docx', ContentFile(content)


def get_soffice_binary():
    configured = getattr(settings, 'LIBREOFFICE_BINARY', '')
    candidates = [configured] if configured else []
    candidates.extend(['soffice', 'libreoffice'])
    for candidate in candidates:
        if candidate and shutil.which(candidate):
            return candidate
    return None


def write_generated_docx_to_path(document, path):
    document.generated_file.open('rb')
    try:
        path.write_bytes(document.generated_file.read())
    finally:
        document.generated_file.close()


def convert_docx_to_pdf_path(document, workdir):
    soffice = get_soffice_binary()
    if not soffice:
        raise ValueError(
            'На сервере не найден LibreOffice для конвертации DOCX в PDF. '
            'Установите libreoffice и шрифты fonts-dejavu/fonts-liberation.'
        )

    source_path = workdir / 'source.docx'
    libreoffice_profile = workdir / 'lo_profile'
    libreoffice_profile.mkdir(parents=True, exist_ok=True)
    write_generated_docx_to_path(document, source_path)
    command = [
        soffice,
        '--headless',
        '--nologo',
        '--nofirststartwizard',
        f'-env:UserInstallation={libreoffice_profile.as_uri()}',
        '--convert-to',
        'pdf:writer_pdf_Export',
        '--outdir',
        str(workdir),
        str(source_path),
    ]
    completed = subprocess.run(command, capture_output=True, text=True, timeout=90)
    if completed.returncode != 0:
        error = completed.stderr.strip() or completed.stdout.strip() or 'LibreOffice не смог создать PDF.'
        raise ValueError(f'Ошибка конвертации DOCX в PDF: {error}')

    pdf_path = workdir / 'source.pdf'
    if not pdf_path.exists() or pdf_path.stat().st_size == 0:
        raise ValueError('LibreOffice завершился без ошибки, но PDF-файл не был создан.')
    return pdf_path


def extract_docx_lines(file_field):
    if not file_field:
        return []
    try:
        file_field.open('rb')
        raw = file_field.read()
    finally:
        try:
            file_field.close()
        except Exception:
            pass

    try:
        doc = DocxDocument(io.BytesIO(raw))
    except Exception:
        return []

    lines = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(' | '.join(cells))

    return lines


def parse_decimal_option(value, default=None):
    if value in (None, ''):
        return default
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def insert_wrapped_lines(pdf, title, lines, *, footer=''):
    page = pdf.new_page(width=595, height=842)
    margin = 42
    y = 36
    targets = {}
    page.insert_textbox(
        fitz.Rect(margin, y, 553, y + 42),
        title,
        **fitz_text_kwargs(15),
        color=(0.07, 0.13, 0.11),
        align=1,
    )
    y += 66
    if not lines:
        lines = ['DOCX-файл создан, но текст для предпросмотра извлечь не удалось. Скачайте DOCX без печати для ручной проверки.']

    for line in lines:
        text = str(line or '').strip()
        if not text:
            continue
        estimated_height = max(24, 16 * (len(text) // 86 + 1))
        if y + estimated_height > 760:
            if footer:
                page.insert_textbox(
                    fitz.Rect(margin, 780, 553, 812),
                    footer,
                    **fitz_text_kwargs(8),
                    color=(0.42, 0.47, 0.44),
                    align=1,
                )
            page = pdf.new_page(width=595, height=842)
            y = 42
        if 'исполнитель' in text.lower():
            targets['executor'] = {
                'page_index': pdf.page_count - 1,
                'x': margin,
                'y': y,
                'height': estimated_height,
            }
        page.insert_textbox(
            fitz.Rect(margin, y, 553, y + estimated_height),
            text,
            **fitz_text_kwargs(10.5),
            color=(0.12, 0.16, 0.14),
        )
        y += estimated_height + 8
    if footer:
        page.insert_textbox(
            fitz.Rect(margin, 780, 553, 812),
            footer,
            **fitz_text_kwargs(8),
            color=(0.42, 0.47, 0.44),
            align=1,
        )
    return targets


def apply_stamp_and_watermark(pdf, rule, stamp_options=None, targets=None):
    stamp_options = stamp_options or {}
    targets = targets or {}
    for page in pdf:
        if rule.watermark_enabled:
            watermark_rect = watermark_rect_for_rule(rule, page.rect)
            if rule.watermark_image:
                page.insert_image(watermark_rect, filename=rule.watermark_image.path, keep_proportion=True, overlay=False)
            elif rule.watermark_text:
                page.insert_textbox(
                    watermark_rect,
                    rule.watermark_text,
                    **fitz_text_kwargs(max(12, min(42, watermark_rect.width / 9))),
                    color=(0.78, 0.82, 0.80),
                    align=1,
                    overlay=False,
                )

    if pdf.page_count:
        page_index, stamp_rect = stamp_rect_for_options(rule, pdf, stamp_options, targets)
        page = pdf[page_index]
        page.insert_image(stamp_rect, filename=rule.stamp_image.path, keep_proportion=True, overlay=True)


def build_fast_stamp_preview_file(document, stamp_options=None):
    rule = find_stamp_rule(document)
    if not rule or not rule.stamp_image:
        raise ValueError('Для этого шаблона не настроена электронная печать.')

    pdf = fitz.open()
    try:
        lines = extract_docx_lines(document.generated_file)
        footer = f'{document.display_title} · предварительный просмотр печати'
        targets = insert_wrapped_lines(pdf, document.display_title, lines, footer=footer)
        if pdf.page_count == 0:
            pdf.new_page(width=595, height=842)
        apply_stamp_and_watermark(pdf, rule, stamp_options=stamp_options, targets=targets)
        buffer = io.BytesIO()
        pdf.save(buffer, deflate=True, garbage=3)
    finally:
        pdf.close()
    buffer.seek(0)

    base = safe_file_slug(document.display_title, 'stamp-preview')
    return f'{base}-stamp-preview-{uuid4().hex[:10]}.pdf', ContentFile(buffer.getvalue())


def build_approved_document_file(document, with_stamp=False, stamp_options=None):
    if not with_stamp:
        return copy_generated_file(document)

    rule = find_stamp_rule(document)
    if not rule or not rule.stamp_image:
        raise ValueError('Для этого шаблона не настроена электронная печать.')

    try:
        with tempfile.TemporaryDirectory(prefix='erp_document_pdf_') as tmp_dir:
            workdir = Path(tmp_dir)
            source_pdf = convert_docx_to_pdf_path(document, workdir)
            try:
                pdf = fitz.open(str(source_pdf))
            except Exception as exc:
                raise ValueError(f'PDF создан, но его не удалось открыть: {exc}') from exc
            try:
                if pdf.page_count == 0:
                    raise ValueError('PDF создан без страниц. Проверьте DOCX-шаблон.')
                targets = find_pdf_text_targets(pdf)
                apply_stamp_and_watermark(pdf, rule, stamp_options=stamp_options, targets=targets)
                output_pdf = workdir / 'approved.pdf'
                pdf.save(str(output_pdf), deflate=True, garbage=3)
                buffer = io.BytesIO(output_pdf.read_bytes())
            finally:
                pdf.close()
    except subprocess.TimeoutExpired as exc:
        raise ValueError('Конвертация DOCX в PDF заняла слишком много времени и была остановлена.') from exc
    buffer.seek(0)

    base = safe_file_slug(document.display_title, 'approved-document')
    return f'{base}-stamped-{uuid4().hex[:10]}.pdf', ContentFile(buffer.getvalue())


def stamp_size_for_options(rule, stamp_options=None):
    stamp_options = stamp_options or {}
    width_mm = parse_decimal_option(stamp_options.get('stamp_width_mm'), rule.width_mm)
    height_mm = parse_decimal_option(stamp_options.get('stamp_height_mm'), rule.height_mm or width_mm)
    width_mm = max(10, min(120, width_mm or rule.width_mm or 40))
    height_mm = max(10, min(120, height_mm or rule.height_mm or width_mm))
    return mm_to_pt(width_mm), mm_to_pt(height_mm)


def mm_to_pt(value):
    return float(value or 0) * 72 / 25.4


def clamp_rect(x, y, width, height, page_rect):
    x = max(0, min(float(x), page_rect.width - width))
    y = max(0, min(float(y), page_rect.height - height))
    return fitz.Rect(x, y, x + width, y + height)


def stamp_rect_for_options(rule, pdf, stamp_options=None, targets=None):
    stamp_options = stamp_options or {}
    targets = targets or {}
    mode = stamp_options.get('stamp_mode') or 'executor'
    requested_position = stamp_options.get('stamp_position') or stamp_options.get('position') or ''
    width, height = stamp_size_for_options(rule, stamp_options)
    raw_page_number = parse_decimal_option(stamp_options.get('page_number'), None)
    if raw_page_number is None:
        raw_page_number = parse_decimal_option(getattr(rule, 'page_number', None), None)
    if raw_page_number is not None:
        page_index = int(max(1, min(pdf.page_count, raw_page_number))) - 1
    else:
        page_index = max(0, pdf.page_count - 1)
    page_rect = pdf[page_index].rect

    if requested_position in {
        StampRule.POSITION_BOTTOM_LEFT,
        StampRule.POSITION_BOTTOM_RIGHT,
        StampRule.POSITION_TOP_LEFT,
        StampRule.POSITION_TOP_RIGHT,
        StampRule.POSITION_CENTER,
    }:
        return page_index, stamp_rect_for_position(requested_position, page_rect, width=width, height=height)

    if mode == 'manual':
        x_percent = parse_decimal_option(stamp_options.get('stamp_x_percent'))
        y_percent = parse_decimal_option(stamp_options.get('stamp_y_percent'))
        if x_percent is not None and y_percent is not None:
            x = page_rect.width * max(0, min(100, x_percent)) / 100
            y = page_rect.height * max(0, min(100, y_percent)) / 100
            return page_index, clamp_rect(x, y, width, height, page_rect)
        x_mm = parse_decimal_option(stamp_options.get('stamp_x_mm'))
        y_mm = parse_decimal_option(stamp_options.get('stamp_y_mm'))
        if x_mm is not None and y_mm is not None:
            return page_index, clamp_rect(mm_to_pt(x_mm), mm_to_pt(y_mm), width, height, page_rect)

    if mode == 'rule':
        return page_index, stamp_rect_for_rule(rule, page_rect, width=width, height=height)

    executor = targets.get('executor')
    if executor:
        page_index = max(0, min(pdf.page_count - 1, int(executor.get('page_index') or 0)))
        page_rect = pdf[page_index].rect
        x = float(executor.get('x') or mm_to_pt(18))
        y = float(executor.get('y') or page_rect.height - mm_to_pt(50)) - height * 0.45
        return page_index, clamp_rect(x, y, width, height, page_rect)

    return page_index, stamp_rect_for_rule(rule, page_rect, width=width, height=height)


def stamp_rect_for_position(position, page_rect, *, width, height):
    margin = mm_to_pt(18)
    if position == StampRule.POSITION_BOTTOM_RIGHT:
        return fitz.Rect(page_rect.width - margin - width, page_rect.height - margin - height, page_rect.width - margin, page_rect.height - margin)
    if position == StampRule.POSITION_TOP_LEFT:
        return fitz.Rect(margin, margin, margin + width, margin + height)
    if position == StampRule.POSITION_TOP_RIGHT:
        return fitz.Rect(page_rect.width - margin - width, margin, page_rect.width - margin, margin + height)
    if position == StampRule.POSITION_CENTER:
        x = (page_rect.width - width) / 2
        y = (page_rect.height - height) / 2
        return fitz.Rect(x, y, x + width, y + height)
    return fitz.Rect(margin, page_rect.height - margin - height, margin + width, page_rect.height - margin)


def find_pdf_text_targets(pdf):
    targets = {}
    for page_index, page in enumerate(pdf):
        matches = page.search_for('Исполнитель')
        if matches:
            rect = matches[-1]
            targets['executor'] = {
                'page_index': page_index,
                'x': rect.x0,
                'y': rect.y0,
                'height': rect.height,
            }
            continue
        try:
            words = page.get_text('words') or []
        except Exception:
            words = []
        for word in words:
            text = str(word[4] if len(word) > 4 else '')
            if 'исполнитель' in text.lower():
                targets['executor'] = {
                    'page_index': page_index,
                    'x': word[0],
                    'y': word[1],
                    'height': max(8, word[3] - word[1]),
                }
    return targets


def stamp_rect_for_rule(rule, page_rect, *, width=None, height=None):
    width = width if width is not None else mm_to_pt(rule.width_mm)
    height = height if height is not None else mm_to_pt(rule.height_mm or rule.width_mm)
    margin = mm_to_pt(18)
    if rule.position == StampRule.POSITION_CUSTOM and rule.x_mm is not None and rule.y_mm is not None:
        x = mm_to_pt(rule.x_mm)
        y = mm_to_pt(rule.y_mm)
        return clamp_rect(x, y, width, height, page_rect)
    if rule.position == StampRule.POSITION_BOTTOM_RIGHT:
        return fitz.Rect(page_rect.width - margin - width, page_rect.height - margin - height, page_rect.width - margin, page_rect.height - margin)
    if rule.position == StampRule.POSITION_TOP_LEFT:
        return fitz.Rect(margin, margin, margin + width, margin + height)
    if rule.position == StampRule.POSITION_TOP_RIGHT:
        return fitz.Rect(page_rect.width - margin - width, margin, page_rect.width - margin, margin + height)
    if rule.position == StampRule.POSITION_CENTER:
        x = (page_rect.width - width) / 2
        y = (page_rect.height - height) / 2
        return fitz.Rect(x, y, x + width, y + height)
    return fitz.Rect(margin, page_rect.height - margin - height, margin + width, page_rect.height - margin)


def watermark_rect_for_rule(rule, page_rect):
    width = mm_to_pt(rule.watermark_width_mm)
    height = mm_to_pt(rule.watermark_height_mm)
    margin = mm_to_pt(18)
    if rule.watermark_position == StampRule.POSITION_CUSTOM and rule.x_mm is not None and rule.y_mm is not None:
        x = mm_to_pt(rule.x_mm)
        y = mm_to_pt(rule.y_mm)
        return fitz.Rect(x, y, x + width, y + height)
    if rule.watermark_position == StampRule.POSITION_BOTTOM_LEFT:
        return fitz.Rect(margin, page_rect.height - margin - height, margin + width, page_rect.height - margin)
    if rule.watermark_position == StampRule.POSITION_BOTTOM_RIGHT:
        return fitz.Rect(page_rect.width - margin - width, page_rect.height - margin - height, page_rect.width - margin, page_rect.height - margin)
    if rule.watermark_position == StampRule.POSITION_TOP_LEFT:
        return fitz.Rect(margin, margin, margin + width, margin + height)
    if rule.watermark_position == StampRule.POSITION_TOP_RIGHT:
        return fitz.Rect(page_rect.width - margin - width, margin, page_rect.width - margin, margin + height)
    x = (page_rect.width - width) / 2
    y = (page_rect.height - height) / 2
    return fitz.Rect(x, y, x + width, y + height)
