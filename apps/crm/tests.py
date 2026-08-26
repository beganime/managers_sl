import tempfile

from django.test import TestCase, override_settings
from docx import Document

from apps.organizations.models import Company
from users.models import User

from .models import Client, ClientQuestionnaire


class ClientQuestionnaireDocumentTests(TestCase):
    def setUp(self):
        self.manager = User.objects.create_user(
            email='questionnaire-manager@example.com',
            password='test-password',
        )
        self.company = Company.objects.create(name='Students Life')
        self.client_record = Client.objects.create(
            company=self.company,
            manager=self.manager,
            full_name='Иванов Иван Иванович',
            phone='+99360000000',
            academic_year=2027,
            funding_type='government',
        )

    def document_text(self, path):
        document = Document(path)
        values = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                values.extend(cell.text for cell in row.cells)
        return '\n'.join(values)

    def test_generated_questionnaire_uses_russian_labels_and_new_fields(self):
        questionnaire = ClientQuestionnaire.objects.create(
            client=self.client_record,
            full_name=self.client_record.full_name,
            phone=self.client_record.phone,
            status=ClientQuestionnaire.STATUS_APPROVED,
            data={
                'form_type': 'applicant',
                'full_name': self.client_record.full_name,
                'is_conscript': True,
                'passport_pending': True,
                'academic_year': 2027,
                'funding_type': 'government',
                'requested_services': ['Подбор вуза', 'Гослиния'],
                'request_text': 'Хочу поступить на лечебное дело.',
                'desired_universities': 'РУДН / БГМУ',
                'desired_program': 'Лечебное дело',
                'university_choices': [
                    {
                        'university_name': 'БГМУ',
                        'programs': [{'name': 'Лечебное дело'}],
                    }
                ],
            },
        )

        with tempfile.TemporaryDirectory() as media_root, override_settings(MEDIA_ROOT=media_root):
            questionnaire.generate_file()
            text = self.document_text(questionnaire.generated_file.path)

        for expected in (
            'Призывник', 'Оформляется: Да',
            'Год поступления', 'Форма поступления', 'Гослиния',
            'Нужные услуги', 'Что хочет клиент', 'Желаемые вузы',
            'Выбранные вузы и программы', 'БГМУ', 'Лечебное дело',
        ):
            self.assertIn(expected, text)
        for api_name in (
            'is_conscript', 'passport_pending', 'academic_year',
            'funding_type', 'requested_services', 'request_text',
            'desired_universities', 'university_choices',
        ):
            self.assertNotIn(api_name, text)
