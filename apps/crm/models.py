import io
import os
from uuid import uuid4

from django.conf import settings
from django.core.files.base import ContentFile
from django.db import models
from django.utils import timezone

from apps.core.models import ActiveModel, TimeStampedModel
from apps.employees.models import EmployeeProfile
from apps.organizations.models import Company, Office


def client_questionnaire_document_upload_to(instance, filename):
    return f'erp/crm/client_questionnaires/{instance.client_id}/{uuid4().hex}-{filename}'


class LeadSource(TimeStampedModel, ActiveModel):
    name = models.CharField('Название источника', max_length=150, unique=True)
    code = models.SlugField('Код источника', max_length=80, unique=True)
    description = models.TextField('Описание', blank=True)

    class Meta:
        verbose_name = 'Источник лида'
        verbose_name_plural = 'Источники лидов'
        ordering = ['name']

    def __str__(self):
        return self.name


class Lead(TimeStampedModel):
    STATUS_CHOICES = (
        ('new', 'Новый'),
        ('contacted', 'Связались'),
        ('qualified', 'Квалифицирован'),
        ('converted', 'Стал клиентом'),
        ('lost', 'Потерян'),
        ('spam', 'Спам'),
    )

    DIRECTION_CHOICES = (
        ('admission', 'Поступление'),
        ('visa', 'Виза'),
        ('translation', 'Переводы'),
        ('tickets', 'Билеты'),
        ('work_visa', 'Рабочие визы'),
        ('other', 'Другое'),
    )

    company = models.ForeignKey(Company, verbose_name='Компания', on_delete=models.PROTECT, related_name='crm_leads')
    office = models.ForeignKey(Office, verbose_name='Офис', on_delete=models.SET_NULL, null=True, blank=True, related_name='crm_leads')
    source = models.ForeignKey(LeadSource, verbose_name='Источник', on_delete=models.SET_NULL, null=True, blank=True, related_name='leads')
    manager = models.ForeignKey(settings.AUTH_USER_MODEL, verbose_name='Ответственный менеджер', on_delete=models.SET_NULL, null=True, blank=True, related_name='crm_leads')

    full_name = models.CharField('ФИО / Имя', max_length=255)
    phone = models.CharField('Телефон', max_length=50, db_index=True)
    email = models.EmailField('Email', blank=True, null=True)
    country = models.CharField('Страна', max_length=100, blank=True)
    city = models.CharField('Город', max_length=100, blank=True)
    direction = models.CharField('Направление', max_length=50, choices=DIRECTION_CHOICES, blank=True)
    interested_country = models.CharField('Интересующая страна обучения', max_length=100, blank=True)
    interested_program = models.CharField('Интересующая программа', max_length=255, blank=True)
    status = models.CharField('Статус', max_length=32, choices=STATUS_CHOICES, default='new', db_index=True)
    comment = models.TextField('Комментарий', blank=True)
    custom_data = models.JSONField('Дополнительные данные', default=dict, blank=True)
    submitter_ip = models.GenericIPAddressField('IP отправителя', null=True, blank=True, db_index=True)
    submitter_user_agent = models.TextField('User-Agent', blank=True, default='')
    submitter_referer = models.URLField('Referer', max_length=1000, blank=True, default='')
    submitter_origin = models.URLField('Origin', max_length=1000, blank=True, default='')
    api_source = models.CharField('API source', max_length=100, blank=True, default='')
    taken_at = models.DateTimeField('Дата взятия ответственности', null=True, blank=True)
    converted_at = models.DateTimeField('Дата конвертации', null=True, blank=True)
    is_archived = models.BooleanField('В архиве', default=False, db_index=True)
    archived_at = models.DateTimeField('Дата архивации', null=True, blank=True)
    archived_by = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        verbose_name='Кто архивировал',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='archived_crm_leads',
    )
    archive_reason = models.TextField('Причина архивации', blank=True)

    class Meta:
        verbose_name = 'Лид'
        verbose_name_plural = 'Лиды'
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['company', 'office', 'status']),
            models.Index(fields=['manager', 'status']),
            models.Index(fields=['is_archived', 'status']),
            models.Index(fields=['phone']),
        ]

    def __str__(self):
        return f'{self.full_name} ({self.phone})'

    @property
    def action_history(self):
        history = (self.custom_data or {}).get('action_history', [])
        return history if isinstance(history, list) else []

    def log_action(self, action, user=None, note='', save=False):
        data = dict(self.custom_data or {})
        history = list(data.get('action_history') or [])
        history.append({
            'action': action,
            'at': timezone.now().isoformat(),
            'user_id': getattr(user, 'id', None),
            'user': user.get_full_name() or getattr(user, 'email', '') if user else '',
            'note': note or '',
        })
        data['action_history'] = history[-100:]
        self.custom_data = data
        if save:
            self.save(update_fields=['custom_data', 'updated_at'])

    def take_responsibility(self, user, company=None, office=None):
        self.manager = user
        if company is not None:
            self.company = company
        if office is not None:
            self.office = office
        self.status = 'contacted'
        if not self.taken_at:
            self.taken_at = timezone.now()
        self.log_action('take_responsibility', user)
        self.save(update_fields=['manager', 'company', 'office', 'status', 'taken_at', 'custom_data', 'updated_at'])

    def release_responsibility(self, user, note=''):
        self.manager = None
        self.status = 'new'
        self.taken_at = None
        self.log_action('release_responsibility', user, note=note)
        self.save(update_fields=['manager', 'status', 'taken_at', 'custom_data', 'updated_at'])

    def archive(self, user=None, reason=''):
        self.is_archived = True
        self.archived_at = timezone.now()
        self.archived_by = user
        self.archive_reason = reason or ''
        self.log_action('archive', user, note=reason)
        self.save(update_fields=['is_archived', 'archived_at', 'archived_by', 'archive_reason', 'custom_data', 'updated_at'])

    def restore_from_archive(self, user=None, note=''):
        self.is_archived = False
        self.archived_at = None
        self.archived_by = None
        self.archive_reason = ''
        self.log_action('restore', user, note=note)
        self.save(update_fields=['is_archived', 'archived_at', 'archived_by', 'archive_reason', 'custom_data', 'updated_at'])

    def mark_converted(self, user=None):
        self.status = 'converted'
        self.converted_at = timezone.now()
        self.log_action('convert_to_client', user)
        self.save(update_fields=['manager', 'company', 'office', 'status', 'converted_at', 'custom_data', 'updated_at'])


class Client(TimeStampedModel):
    STATUS_CHOICES = (
        ('new', 'Новый'),
        ('consultation', 'Консультация'),
        ('documents', 'Сбор документов'),
        ('application', 'Подача заявки'),
        ('invitation', 'Приглашение'),
        ('visa', 'Виза'),
        ('arrived', 'Прибыл'),
        ('success', 'Завершён успешно'),
        ('rejected', 'Отказ'),
        ('archive', 'Архив'),
    )

    company = models.ForeignKey(Company, verbose_name='Компания', on_delete=models.PROTECT, related_name='crm_clients')
    office = models.ForeignKey(Office, verbose_name='Офис', on_delete=models.SET_NULL, null=True, blank=True, related_name='crm_clients')
    manager = models.ForeignKey(settings.AUTH_USER_MODEL, verbose_name='Основной менеджер', on_delete=models.PROTECT, related_name='crm_clients')
    shared_with = models.ManyToManyField(settings.AUTH_USER_MODEL, blank=True, related_name='shared_crm_clients', verbose_name='Доступ открыт также для')
    source_lead = models.OneToOneField(Lead, verbose_name='Исходный лид', on_delete=models.SET_NULL, null=True, blank=True, related_name='client')
    lead_source = models.ForeignKey(LeadSource, verbose_name='Источник', on_delete=models.SET_NULL, null=True, blank=True, related_name='crm_clients')
    direction = models.CharField('Направление', max_length=50, choices=Lead.DIRECTION_CHOICES, blank=True)

    full_name = models.CharField('ФИО клиента', max_length=255, db_index=True)
    phone = models.CharField('Телефон', max_length=50, db_index=True)
    email = models.EmailField('Email', blank=True, null=True)
    dob = models.DateField('Дата рождения', null=True, blank=True)
    citizenship = models.CharField('Гражданство', max_length=100, blank=True)
    city = models.CharField('Город', max_length=100, blank=True)
    address = models.TextField('Адрес проживания', blank=True)
    address_registration = models.TextField('Адрес регистрации', blank=True)
    passport_local_num = models.CharField('Внутренний паспорт', max_length=50, blank=True)
    passport_inter_num = models.CharField('Загранпаспорт', max_length=50, blank=True)
    passport_issued_by = models.CharField('Кем выдан паспорт', max_length=255, blank=True)
    passport_issued_date = models.DateField('Дата выдачи паспорта', null=True, blank=True)
    passport_valid_until = models.DateField('Срок действия паспорта', null=True, blank=True)
    passport_birth_place = models.CharField('Место рождения', max_length=255, blank=True)
    relative_full_name = models.CharField('ФИО родственника', max_length=255, blank=True)
    relative_relation = models.CharField('Кем приходится', max_length=120, blank=True)
    relative_phone = models.CharField('Телефон родственника', max_length=80, blank=True)
    relative_workplace = models.CharField('Место работы родственника', max_length=255, blank=True)
    current_education = models.CharField('Текущее образование', max_length=255, blank=True)
    current_school = models.CharField('Текущий вуз / школа', max_length=255, blank=True)
    current_study_country = models.CharField('Страна текущего обучения', max_length=100, blank=True)
    interested_country = models.CharField('Интересующая страна', max_length=100, blank=True)
    interested_university = models.CharField('Интересующий вуз', max_length=255, blank=True)
    interested_program = models.CharField('Интересующая программа', max_length=255, blank=True)
    has_passport = models.BooleanField('Есть паспорт', default=False)
    has_education_doc = models.BooleanField('Есть аттестат / диплом', default=False)
    has_translation = models.BooleanField('Есть перевод', default=False)
    has_photo = models.BooleanField('Есть фото', default=False)
    status = models.CharField('Статус', max_length=32, choices=STATUS_CHOICES, default='new', db_index=True)
    is_priority = models.BooleanField('Приоритетный клиент', default=False)
    is_partner_client = models.BooleanField('Клиент от партнёра', default=False)
    partner_name = models.CharField('Партнёр', max_length=255, blank=True)
    comments = models.TextField('Комментарии', blank=True)
    custom_data = models.JSONField('Дополнительные данные', default=dict, blank=True)

    class Meta:
        verbose_name = 'Клиент CRM'
        verbose_name_plural = 'Клиенты CRM'
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['company', 'office', 'status']),
            models.Index(fields=['manager', 'status']),
            models.Index(fields=['phone']),
            models.Index(fields=['full_name']),
        ]

    def __str__(self):
        return f'{self.full_name} [{self.get_status_display()}]'


    mobile_app_user_id = models.PositiveIntegerField('Mobile app user ID', null=True, blank=True, db_index=True)
    mobile_app_source = models.BooleanField('Mobile app client', default=False)


class Application(TimeStampedModel):
    STATUS_CHOICES = (
        ('draft', 'Черновик'),
        ('documents', 'Сбор документов'),
        ('submitted', 'Подано'),
        ('in_review', 'На рассмотрении'),
        ('accepted', 'Принято'),
        ('invitation', 'Приглашение получено'),
        ('visa', 'Виза'),
        ('enrolled', 'Зачислен'),
        ('rejected', 'Отказ'),
        ('cancelled', 'Отменено'),
    )

    client = models.ForeignKey(Client, verbose_name='Клиент', on_delete=models.CASCADE, related_name='applications')
    company = models.ForeignKey(Company, verbose_name='Компания', on_delete=models.PROTECT, related_name='crm_applications')
    office = models.ForeignKey(Office, verbose_name='Офис', on_delete=models.SET_NULL, null=True, blank=True, related_name='crm_applications')
    manager = models.ForeignKey(settings.AUTH_USER_MODEL, verbose_name='Ответственный', on_delete=models.PROTECT, related_name='crm_applications')

    university_name = models.CharField('ВУЗ', max_length=255, blank=True)
    program_name = models.CharField('Программа', max_length=255, blank=True)
    country = models.CharField('Страна обучения', max_length=100, blank=True)
    degree = models.CharField('Степень', max_length=100, blank=True)
    language = models.CharField('Язык обучения', max_length=100, blank=True)
    intake = models.CharField('Набор / intake', max_length=100, blank=True)
    status = models.CharField('Статус заявки', max_length=32, choices=STATUS_CHOICES, default='draft', db_index=True)
    submitted_at = models.DateField('Дата подачи', null=True, blank=True)
    decision_at = models.DateField('Дата решения', null=True, blank=True)
    comment = models.TextField('Комментарий', blank=True)
    custom_data = models.JSONField('Дополнительные данные', default=dict, blank=True)

    class Meta:
        verbose_name = 'Заявка на поступление'
        verbose_name_plural = 'Заявки на поступление'
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['company', 'office', 'status']),
            models.Index(fields=['manager', 'status']),
            models.Index(fields=['client', 'status']),
        ]

    def __str__(self):
        return f'{self.client.full_name} — {self.university_name or "ВУЗ не выбран"}'


class ClientActivity(TimeStampedModel):
    ACTIVITY_TYPE_CHOICES = (
        ('call', 'Звонок'),
        ('message', 'Сообщение'),
        ('meeting', 'Встреча'),
        ('note', 'Заметка'),
        ('status_change', 'Смена статуса'),
        ('document', 'Документ'),
        ('payment', 'Платёж'),
    )

    client = models.ForeignKey(Client, verbose_name='Клиент', on_delete=models.CASCADE, related_name='activities')
    manager = models.ForeignKey(settings.AUTH_USER_MODEL, verbose_name='Менеджер', on_delete=models.SET_NULL, null=True, blank=True, related_name='crm_activities')
    activity_type = models.CharField('Тип активности', max_length=32, choices=ACTIVITY_TYPE_CHOICES, default='note')
    title = models.CharField('Заголовок', max_length=255)
    description = models.TextField('Описание', blank=True)
    due_at = models.DateTimeField('Срок / напоминание', null=True, blank=True)
    completed_at = models.DateTimeField('Выполнено', null=True, blank=True)

    class Meta:
        verbose_name = 'Активность клиента'
        verbose_name_plural = 'Активности клиентов'
        ordering = ['-created_at']

    def __str__(self):
        return f'{self.get_activity_type_display()}: {self.client.full_name}'


class ClientNote(TimeStampedModel):
    client = models.ForeignKey(Client, verbose_name='Клиент', on_delete=models.CASCADE, related_name='notes')
    author = models.ForeignKey(settings.AUTH_USER_MODEL, verbose_name='Автор', on_delete=models.SET_NULL, null=True, blank=True, related_name='crm_notes')
    text = models.TextField('Текст заметки')
    is_private = models.BooleanField('Приватная заметка', default=False)

    class Meta:
        verbose_name = 'Заметка клиента'
        verbose_name_plural = 'Заметки клиентов'
        ordering = ['-created_at']

    def __str__(self):
        return f'Заметка: {self.client.full_name}'


class ClientFile(TimeStampedModel):
    STATUS_PENDING = 'pending'
    STATUS_APPROVED = 'approved'
    STATUS_REJECTED = 'rejected'
    STATUS_CHOICES = (
        (STATUS_PENDING, 'Pending review'),
        (STATUS_APPROVED, 'Approved'),
        (STATUS_REJECTED, 'Rejected'),
    )

    client = models.ForeignKey(Client, verbose_name='Клиент', on_delete=models.CASCADE, related_name='files')
    application = models.ForeignKey(Application, verbose_name='Заявка', on_delete=models.CASCADE, null=True, blank=True, related_name='files')
    uploaded_by = models.ForeignKey(settings.AUTH_USER_MODEL, verbose_name='Кто загрузил', on_delete=models.SET_NULL, null=True, blank=True, related_name='crm_files')
    title = models.CharField('Название файла', max_length=255)
    file = models.FileField('Файл', upload_to='erp/crm/client_files/')
    file_type = models.CharField('Тип файла', max_length=100, blank=True)
    comment = models.TextField('Комментарий', blank=True)

    class Meta:
        verbose_name = 'Файл клиента'
        verbose_name_plural = 'Файлы клиентов'
        ordering = ['-created_at']

    def __str__(self):
        return self.title

    external_file_url = models.URLField('External file URL', max_length=1000, blank=True)
    external_mobile_document_id = models.PositiveIntegerField('Mobile document ID', null=True, blank=True, db_index=True)
    external_mobile_user_id = models.PositiveIntegerField('Mobile user ID', null=True, blank=True, db_index=True)
    source = models.CharField('Source', max_length=80, blank=True)
    has_translation = models.BooleanField('Has translation', default=False)
    status = models.CharField('Review status', max_length=20, choices=STATUS_CHOICES, default=STATUS_PENDING, db_index=True)
    review_comment = models.TextField('Review comment', blank=True)
    reviewed_at = models.DateTimeField('Reviewed at', null=True, blank=True)
    reviewed_by = models.ForeignKey(settings.AUTH_USER_MODEL, verbose_name='Reviewed by', on_delete=models.SET_NULL, null=True, blank=True, related_name='reviewed_crm_files')
    external_review_data = models.JSONField('Ответ проверки Student’s Life', default=dict, blank=True)

    @property
    def external_reviewed_by_name(self):
        data = self.external_review_data or {}
        return data.get('reviewed_by_name') or ''

    @property
    def external_reviewed_by_email(self):
        data = self.external_review_data or {}
        return data.get('reviewed_by_email') or ''

    @property
    def reviewed_by_display(self):
        data = self.external_review_data or {}
        external_display = data.get('reviewed_by_display') or data.get('reviewed_by_name') or data.get('reviewed_by_email')
        if external_display:
            return external_display
        if self.reviewed_by_id:
            return self.reviewed_by.get_full_name() or self.reviewed_by.email
        return ''


class ManagerDocumentPlan(TimeStampedModel, ActiveModel):
    PERIOD_DAY = 'day'
    PERIOD_WEEK = 'week'
    PERIOD_MONTH = 'month'
    PERIOD_CUSTOM = 'custom'
    PERIOD_CHOICES = (
        (PERIOD_DAY, 'День'),
        (PERIOD_WEEK, 'Неделя'),
        (PERIOD_MONTH, 'Месяц'),
        (PERIOD_CUSTOM, 'Произвольный период'),
    )

    employee = models.ForeignKey(EmployeeProfile, verbose_name='Менеджер', on_delete=models.CASCADE, related_name='document_plans')
    period_type = models.CharField('Период', max_length=16, choices=PERIOD_CHOICES, default=PERIOD_MONTH)
    start_date = models.DateField('Дата начала', db_index=True)
    end_date = models.DateField('Дата окончания', db_index=True)
    target_clients = models.PositiveIntegerField('План по загруженным клиентам', default=0)
    admin_comment = models.TextField('Комментарий администратора', blank=True)

    class Meta:
        verbose_name = 'План менеджера по документам'
        verbose_name_plural = 'Планы менеджеров по документам'
        ordering = ['-start_date', 'employee__user__first_name']
        indexes = [
            models.Index(fields=['employee', 'start_date', 'end_date', 'is_active'], name='crm_mdocplan_emp_period_idx'),
        ]

    def __str__(self):
        return f'{self.employee} — {self.start_date:%d.%m.%Y}-{self.end_date:%d.%m.%Y}'


class ManagerDocumentCredit(TimeStampedModel):
    EVENT_UPLOADED_CLIENT_DOCUMENTS = 'uploaded_client_documents'
    EVENT_CHOICES = (
        (EVENT_UPLOADED_CLIENT_DOCUMENTS, 'Документы клиента загружены'),
    )

    employee = models.ForeignKey(EmployeeProfile, verbose_name='Менеджер', on_delete=models.CASCADE, related_name='document_credits')
    client = models.ForeignKey(Client, verbose_name='Клиент', on_delete=models.CASCADE, related_name='document_manager_credits')
    plan = models.ForeignKey(ManagerDocumentPlan, verbose_name='План', on_delete=models.SET_NULL, null=True, blank=True, related_name='credits')
    event_type = models.CharField('Тип события', max_length=64, choices=EVENT_CHOICES, default=EVENT_UPLOADED_CLIENT_DOCUMENTS, db_index=True)
    period_start = models.DateField('Начало периода', db_index=True)
    period_end = models.DateField('Конец периода', db_index=True)
    credited_by = models.ForeignKey(settings.AUTH_USER_MODEL, verbose_name='Кто засчитал', on_delete=models.SET_NULL, null=True, blank=True, related_name='document_credits_created')
    credited_at = models.DateTimeField('Дата зачёта', default=timezone.now, db_index=True)
    comment = models.TextField('Комментарий', blank=True)

    class Meta:
        verbose_name = 'Зачёт загруженного клиента'
        verbose_name_plural = 'Зачёты загруженных клиентов'
        ordering = ['-credited_at']
        constraints = [
            models.UniqueConstraint(
                fields=['employee', 'client', 'event_type', 'period_start', 'period_end'],
                name='uniq_manager_client_document_credit_period',
            ),
        ]
        indexes = [
            models.Index(fields=['employee', 'period_start', 'period_end'], name='crm_mdoccredit_emp_period_idx'),
            models.Index(fields=['client', 'event_type'], name='crm_mdoccredit_client_idx'),
        ]

    def __str__(self):
        return f'{self.employee} +1 {self.client} ({self.period_start:%d.%m.%Y})'


class ClientQuestionnaire(TimeStampedModel):
    STATUS_DRAFT = 'draft'
    STATUS_COMPLETED = 'completed'
    STATUS_SUBMITTED = 'submitted'
    STATUS_APPROVED = 'approved'
    STATUS_REJECTED = 'rejected'
    STATUS_UPDATED = 'updated'
    STATUS_CHOICES = (
        (STATUS_DRAFT, 'Не заполнена'),
        (STATUS_COMPLETED, 'Заполнена'),
        (STATUS_SUBMITTED, 'Отправлена на проверку'),
        (STATUS_APPROVED, 'Принята'),
        (STATUS_REJECTED, 'Отклонена'),
        (STATUS_UPDATED, 'Обновлена'),
    )

    client = models.OneToOneField(Client, verbose_name='Клиент', on_delete=models.CASCADE, related_name='questionnaire')
    mobile_questionnaire_id = models.PositiveIntegerField('Mobile questionnaire ID', null=True, blank=True, db_index=True)
    external_mobile_user_id = models.PositiveIntegerField('Mobile user ID', null=True, blank=True, db_index=True)
    source = models.CharField('Источник', max_length=80, blank=True, default='students_life_mobile_app')
    status = models.CharField('Статус анкеты', max_length=20, choices=STATUS_CHOICES, default=STATUS_DRAFT, db_index=True)

    full_name = models.CharField('ФИО', max_length=255, blank=True)
    phone = models.CharField('Телефон', max_length=80, blank=True)
    email = models.EmailField('Email', blank=True, null=True)
    citizenship = models.CharField('Гражданство', max_length=120, blank=True)
    desired_program = models.CharField('Желаемая программа / Вуз', max_length=255, blank=True)
    desired_country = models.CharField('Желаемая страна', max_length=120, blank=True)
    desired_city = models.CharField('Желаемый город', max_length=120, blank=True)
    face_photo_url = models.URLField('Фото абитуриента', max_length=1000, blank=True)
    data = models.JSONField('Данные анкеты', default=dict, blank=True)
    generated_file = models.FileField('Сгенерированный документ', upload_to=client_questionnaire_document_upload_to, blank=True, null=True)
    submitted_at = models.DateTimeField('Дата заполнения', null=True, blank=True)
    last_synced_at = models.DateTimeField('Последняя синхронизация', null=True, blank=True)

    @property
    def reviewed_at_external(self):
        return (self.data or {}).get('reviewed_at') or ''

    @property
    def reviewed_by_display(self):
        data = self.data or {}
        return data.get('reviewed_by_display') or data.get('reviewed_by_name') or data.get('reviewed_by_email') or ''

    @property
    def reviewed_by_name(self):
        return (self.data or {}).get('reviewed_by_name') or ''

    @property
    def reviewed_by_email(self):
        return (self.data or {}).get('reviewed_by_email') or ''

    @property
    def review_comment(self):
        return (self.data or {}).get('review_comment') or (self.data or {}).get('comment') or ''

    class Meta:
        verbose_name = 'Анкета клиента'
        verbose_name_plural = 'Анкеты клиентов'
        ordering = ['-updated_at']
        indexes = [
            models.Index(fields=['status', 'updated_at']),
            models.Index(fields=['external_mobile_user_id']),
        ]

    def __str__(self):
        return f'Анкета: {self.full_name or self.client.full_name}'

    def generate_file(self):
        from docx import Document
        from docx.shared import Inches, Pt

        template_path = os.path.join(os.path.dirname(__file__), 'document_templates', 'anketa_students_life_template_v2.docx')
        document = Document(template_path) if os.path.exists(template_path) else Document()
        data = self.data or {}

        value_labels = {
            'male': '???????',
            'female': '???????',
            'school_student': '???????? / ??????????????? ??????',
            'applicant': '?????????? / ?????? ??????',
            True: '??',
            False: '???',
        }
        labels = {
            'full_name': '???',
            'birth_date': '???? ????????',
            'gender': '???',
            'citizenship': '???????????',
            'marital_status': '???????? ?????????',
            'residence_country': '??????',
            'residence_region': '??????',
            'residence_city': '?????',
            'residence_street': '?????',
            'residence_house': '??? / ????????',
            'residence_postal_code': '???????? ??????',
            'passport_number': '???????',
            'passport_issued_by': '??? ????????',
            'passport_issue_date': '???? ?????? ????????',
            'passport_expiry_date': '???? ????????? ????????',
            'has_international_passport': '?????????????',
            'phone': '???????',
            'email': 'Email',
            'extra_phone': '???. ???????',
            'imo': 'Imo',
            'telegram': 'Telegram',
            'preferred_contact_method': '?????? ?????',
            'parent_full_name': '??? ????????',
            'parent_relation': '??? ????????',
            'parent_contacts': '???????? ????????',
            'parent_workplace': '?????? ????????',
            'family_members': '?????',
            'education_level': '??????? ???????????',
            'school_class': '?????',
            'school_name': '??????? ?????????',
            'school_country': '?????? ?????',
            'school_city': '????? ?????',
            'graduation_year': '??? ?????????',
            'education_status': '??????',
            'desired_program': '????????? / ???',
            'admission_goal': '????',
            'desired_city': '?????',
            'desired_country': '??????',
            'desired_language': '???? ????????',
            'desired_education_level': '??????? ????????',
            'admission_urgency': '?????????',
            'has_visa': '????',
            'visa_country': '?????? ????',
            'visa_city': '????? ????',
            'visa_valid_until': '???? ????',
            'hobbies': '?????',
            'applicant_comment': '???????????',
            'referral_source': '????????',
            'data_processing_consent': '???????? ?? ????????? ??????',
            'status': '?????? ??????',
            'generated_document_at': '???? ????????????',
        }

        from .questionnaire_labels import QUESTIONNAIRE_DOCUMENT_LABELS, QUESTIONNAIRE_VALUE_LABELS

        labels = QUESTIONNAIRE_DOCUMENT_LABELS
        value_labels = QUESTIONNAIRE_VALUE_LABELS

        def render_value(value):
            try:
                if value in value_labels:
                    return value_labels[value]
            except TypeError:
                pass
            if value in (None, '', [], {}):
                return '-'
            if hasattr(value, 'strftime'):
                return timezone.localtime(value).strftime('%d.%m.%Y %H:%M') if hasattr(value, 'tzinfo') and value.tzinfo else value.strftime('%d.%m.%Y')
            if isinstance(value, list):
                items = []
                for item in value:
                    if isinstance(item, dict):
                        language = item.get('language') or item.get('name') or item.get('title')
                        level = item.get('level')
                        items.append(f'{language} - {level}' if language and level else str(language or item))
                    else:
                        items.append(str(item))
                return '\n'.join(f'- {item}' for item in items) if items else '-' 
            if isinstance(value, dict):
                return '\n'.join(f'{labels.get(str(key), key)}: {render_value(val)}' for key, val in value.items())
            return str(value)

        def clear_cell(cell):
            for paragraph in cell.paragraphs:
                paragraph.clear()

        def set_cell(cell, value, bold=False, size=8):
            clear_cell(cell)
            run = cell.paragraphs[0].add_run(str(value or ''))
            run.bold = bold
            run.font.size = Pt(size)

        def set_pair(table, row_index, field_one, value_one, field_two=None, value_two=None):
            cells = table.rows[row_index].cells
            set_cell(cells[0], labels.get(field_one, field_one), bold=True)
            set_cell(cells[3], render_value(value_one))
            if field_two and len(cells) > 6:
                set_cell(cells[5], labels.get(field_two, field_two), bold=True)
                set_cell(cells[6], render_value(value_two))
            elif len(cells) > 6:
                set_cell(cells[5], '')
                set_cell(cells[6], '')

        def append_pair(table, field_one, value_one, field_two=None, value_two=None):
            cells = table.add_row().cells
            if len(cells) >= 7:
                set_cell(cells[0], labels.get(field_one, field_one), bold=True)
                set_cell(cells[3], render_value(value_one))
                if field_two:
                    set_cell(cells[5], labels.get(field_two, field_two), bold=True)
                    set_cell(cells[6], render_value(value_two))

        if len(document.tables) >= 14:
            tables = document.tables
            form_type = data.get('form_type') or data.get('application_type') or 'applicant'
            title = 'ПРЕДВАРИТЕЛЬНАЯ ЗАЯВКА ШКОЛЬНИКА' if form_type == 'school_student' else 'АНКЕТА АБИТУРИЕНТА'
            set_cell(
                tables[0].cell(0, 0),
                f'{title}\nПерсональная карточка для поступления и сопровождения\nДата формирования: {timezone.localtime(timezone.now()):%d.%m.%Y %H:%M}',
                bold=True,
                size=12,
            )
            set_cell(tables[0].cell(0, 1), 'ФОТО\n3 × 4 см\nФото см. по ссылке в карточке', bold=True)

            set_pair(tables[1], 1, 'full_name', data.get('full_name') or self.full_name, 'birth_date', data.get('birth_date'))
            set_pair(tables[1], 2, 'gender', data.get('gender'), 'citizenship', data.get('citizenship') or self.citizenship)
            set_pair(tables[1], 3, 'marital_status', data.get('marital_status'))
            set_pair(tables[2], 1, 'residence_country', data.get('residence_country'), 'residence_region', data.get('residence_region'))
            set_pair(tables[2], 2, 'residence_city', data.get('residence_city'), 'residence_street', data.get('residence_street'))
            set_pair(tables[2], 3, 'residence_house', data.get('residence_house'), 'residence_postal_code', data.get('residence_postal_code'))
            set_pair(tables[3], 1, 'passport_number', data.get('passport_number'), 'passport_issued_by', data.get('passport_issued_by'))
            set_pair(tables[3], 2, 'passport_issue_date', data.get('passport_issue_date'), 'passport_expiry_date', data.get('passport_expiry_date'))
            set_pair(tables[3], 3, 'has_international_passport', data.get('has_international_passport'))
            set_pair(tables[4], 1, 'phone', data.get('phone') or self.phone, 'email', data.get('email') or self.email)
            set_pair(tables[4], 2, 'extra_phone', data.get('extra_phone'), 'imo', data.get('imo'))
            set_pair(tables[4], 3, 'telegram', data.get('telegram'), 'preferred_contact_method', data.get('preferred_contact_method'))
            set_pair(tables[5], 1, 'parent_full_name', data.get('parent_full_name'), 'parent_relation', data.get('parent_relation'))
            set_pair(tables[5], 2, 'parent_contacts', data.get('parent_contacts'), 'parent_workplace', data.get('parent_workplace'))
            set_pair(tables[5], 3, 'family_members', data.get('family_members'))
            set_pair(tables[6], 1, 'education_level', data.get('education_level'), 'school_name', data.get('school_name'))
            set_pair(tables[6], 2, 'school_country', data.get('school_country'), 'school_city', data.get('school_city'))
            set_pair(tables[6], 3, 'graduation_year', data.get('graduation_year'), 'education_status', data.get('education_status'))
            if data.get('school_class'):
                append_pair(tables[6], 'school_class', data.get('school_class'))
            set_pair(tables[7], 1, 'desired_program', data.get('desired_program') or self.desired_program, 'admission_goal', data.get('admission_goal'))
            set_pair(tables[7], 2, 'desired_city', data.get('desired_city') or self.desired_city, 'desired_country', data.get('desired_country') or self.desired_country)
            set_pair(tables[7], 3, 'desired_language', data.get('desired_language'), 'desired_education_level', data.get('desired_education_level'))
            set_pair(tables[7], 4, 'admission_urgency', data.get('admission_urgency'))
            set_pair(tables[8], 1, 'has_visa', data.get('has_visa'))
            if data.get('visa_country') or data.get('visa_city'):
                append_pair(tables[8], 'visa_country', data.get('visa_country'), 'visa_city', data.get('visa_city'))
            if data.get('visa_valid_until'):
                append_pair(tables[8], 'visa_valid_until', data.get('visa_valid_until'))
            set_pair(tables[9], 1, 'referral_source', data.get('referral_source'))
            if data.get('hobbies'):
                append_pair(tables[9], 'hobbies', data.get('hobbies'))
            if data.get('applicant_comment'):
                append_pair(tables[9], 'applicant_comment', data.get('applicant_comment'))
            append_pair(tables[9], 'data_processing_consent', data.get('data_processing_consent'))
            languages = data.get('languages') or [{'language': labels.get('languages', 'Languages'), 'level': '-'}]
            for index, item in enumerate(languages):
                cells = tables[10].rows[index + 1].cells if index + 1 < len(tables[10].rows) else tables[10].add_row().cells
                language = item.get('language') if isinstance(item, dict) else str(item)
                level = item.get('level') if isinstance(item, dict) else '-'
                set_cell(cells[0], language, bold=True)
                set_cell(cells[3], render_value(level))
            set_cell(tables[11].rows[1].cells[0], render_value(data.get('achievements')))
            set_cell(tables[12].rows[1].cells[0], render_value(data.get('help_needed')))
            set_pair(tables[13], 1, 'status', self.get_status_display())
            set_pair(tables[13], 2, 'generated_document_at', timezone.now())
        else:
            section = document.sections[0]
            section.top_margin = Inches(0.47)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            table = document.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            for field in ('full_name', 'phone', 'email', 'citizenship', 'desired_program', 'desired_country', 'desired_city'):
                row = table.add_row().cells
                set_cell(row[0], labels.get(field, field), bold=True)
                set_cell(row[1], render_value(data.get(field) or getattr(self, field, '')))

        buffer = io.BytesIO()
        document.save(buffer)
        filename = f'anketa-{self.client_id}-{uuid4().hex[:8]}.docx'
        self.generated_file.save(filename, ContentFile(buffer.getvalue()), save=False)
        self.save(update_fields=['generated_file', 'updated_at'])
        return self.generated_file
